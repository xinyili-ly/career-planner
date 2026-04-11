import unittest
import requests
import subprocess
import time
import os
import sys
from docx import Document

# 模拟简历文本
MOCK_RESUME = """
姓名：张三
学校：电子科技大学（211）
专业：计算机科学与技术（本科）
毕业时间：2025年

【技能】
- 熟练掌握 Python, Java, C++
- 熟悉 Django, FastAPI 后端框架
- 了解 MySQL, Redis 数据库
- 英语六级 (CET-6 520分)

【经历】
1. 校园二手交易平台（后端负责人）
   - 使用 FastAPI + Vue.js 开发，日活 500+
   - 负责数据库设计与 API 接口开发
2. 某互联网公司 Java 实习生（3个月）
   - 参与营销活动后台开发，负责数据报表导出功能

【意向】
- 期望职位：后端开发、Python开发
- 期望城市：成都、深圳
- 薪资：12k-15k
"""

# 默认 URL，可能会在 setUpClass 中根据端口调整
BASE_URL = "http://127.0.0.1:8000"
# 定位到 backend 目录，以便正确启动 uvicorn
BACKEND_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), '../backend'))

class TestCareerAgentFlow(unittest.TestCase):
    server_process = None
    server_started_by_test = False

    @classmethod
    def setUpClass(cls):
        """
        在所有测试开始前执行：
        1. 检查默认端口 8000 是否有服务在运行
        2. 如果没有，则尝试在 8001 端口启动新服务
        """
        global BASE_URL
        print("\n>>> [Setup] 正在检查后端服务状态...")
        
        # 1. 尝试连接默认端口 8000
        try:
            requests.get("http://127.0.0.1:8000", timeout=1)
            print(">>> 服务已在 8000 端口运行，将直接使用。")
            BASE_URL = "http://127.0.0.1:8000"
            return
        except (requests.exceptions.ConnectionError, requests.exceptions.ReadTimeout):
            print(">>> 默认端口 8000 不可用或无响应。")

        # 2. 尝试启动新服务 (使用端口 8001 避免冲突)
        port = 8001
        BASE_URL = f"http://127.0.0.1:{port}"
        print(f">>> 尝试在端口 {port} 启动新服务...")
        
        # 检查 backend 目录是否存在
        if not os.path.exists(BACKEND_DIR):
            raise RuntimeError(f"后端目录不存在: {BACKEND_DIR}")

        # 启动服务 (非阻塞)
        # 注意：这里我们不捕获 stdout/stderr，以便在控制台能看到 uvicorn 的启动日志（如果有报错的话）
        # 如果觉得日志太吵，可以加上 stdout=subprocess.DEVNULL
        cls.server_process = subprocess.Popen(
            [sys.executable, "-m", "uvicorn", "app.main:app", "--host", "127.0.0.1", "--port", str(port)],
            cwd=BACKEND_DIR
        )
        cls.server_started_by_test = True
        
        # 轮询等待服务启动 (最多等待 30 秒)
        print(">>> 等待服务启动 (Max 30s)...")
        max_retries = 30
        for i in range(max_retries):
            try:
                requests.get(BASE_URL, timeout=3)
                print(f">>> 服务启动成功！(耗时 {i+1}s)")
                break
            except (requests.exceptions.ConnectionError, requests.exceptions.ReadTimeout):
                if i == max_retries - 1:
                    print(">>> 服务启动超时！请检查依赖是否安装完整 (如 fastapi, uvicorn, pdfplumber 等)")
                    if cls.server_process:
                        cls.server_process.terminate()
                    raise RuntimeError("无法启动后端服务")
                time.sleep(1)

    @classmethod
    def tearDownClass(cls):
        """所有测试结束后清理"""
        if cls.server_started_by_test and cls.server_process:
            print("\n>>> [Teardown] 正在关闭测试启动的后端服务...")
            cls.server_process.terminate()
            cls.server_process.wait()
            print(">>> 服务已关闭。")

    def test_01_health_check(self):
        """测试服务健康检查接口"""
        print(f"\n--- Test 01: Health Check ({BASE_URL}) ---")
        try:
            resp = requests.get(f"{BASE_URL}/", timeout=10)
            self.assertEqual(resp.status_code, 200)
            data = resp.json()
            print(f"Server Status: {data}")
            self.assertIn("status", data)
            self.assertEqual(data["status"], "running")
        except Exception as e:
            self.fail(f"健康检查失败: {e}")

    def test_02_manual_parse_flow(self):
        """测试：手动输入文本 -> 解析 -> 榜单 -> 报告"""
        print("\n--- Test 02: Manual Parse Flow ---")
        
        # 1. 解析
        print("[Step 1] Parsing Resume...")
        try:
            # 解析可能极慢，特别是首次下载模型时
            resp = requests.post(
                f"{BASE_URL}/api/v1/student/parse/manual",
                json={"resume_text": MOCK_RESUME},
                timeout=300 # 增加到 5 分钟
            )
            self.assertEqual(resp.status_code, 200, f"解析失败: {resp.text}")
            data_step1 = resp.json()
            
            # 验证结构
            self.assertIn("student_profile", data_step1)
            self.assertIn("leaderboard", data_step1)
            self.assertIn("analysis_summary", data_step1)
            
            profile = data_step1["student_profile"]
            leaderboard = data_step1["leaderboard"]
            
            print(f"Parsed Name: {profile['basic_info']['name']}")
            print(f"Completeness: {data_step1['analysis_summary']['completeness']}%")
            
            # 验证榜单非空
            self.assertTrue(len(leaderboard) > 0, "榜单为空")
            top_job = leaderboard[0]
            print(f"Top Job: {top_job['title']} (Score: {top_job['match_score']})")
            
            # 2. 生成报告
            print(f"[Step 2] Generating Report for {top_job['job_id']}...")
            resp_report = requests.post(
                f"{BASE_URL}/api/v1/match/report",
                json={
                    "student_profile": profile,
                    "job_id": top_job["job_id"]
                },
                timeout=60 # 增加到 60 秒
            )
            self.assertEqual(resp_report.status_code, 200, f"报告生成失败: {resp_report.text}")
            data_step2 = resp_report.json()
            
            # 验证报告内容
            self.assertIn("match_analysis", data_step2)
            self.assertEqual(data_step2["job_info"]["job_id"], top_job["job_id"])
            self.assertIn("career_report", data_step2)
            self.assertIn("module_1", data_step2["career_report"])
            m1 = data_step2["career_report"]["module_1"]
            self.assertIn("profile_overview", m1)
            self.assertIn("target_job_overview", m1)
            self.assertIn("fit_analysis", m1)
            self.assertIn("strengths_and_gaps", m1)
            self.assertIn("overall_conclusion", m1)
            self.assertIn("module_2", data_step2["career_report"])
            m2 = data_step2["career_report"]["module_2"]
            self.assertIsNotNone(m2)
            self.assertIn("goal", m2)
            self.assertIn("market_trend", m2)
            self.assertIn("career_paths", m2)
            cp = m2["career_paths"]
            self.assertIn("path_options", cp)
            self.assertIn("recommended_path", cp)
            self.assertIn("selected_path", cp)
            self.assertIn("module_3", data_step2["career_report"])
            m3 = data_step2["career_report"]["module_3"]
            self.assertIsNotNone(m3)
            self.assertIn("plan_mode", m3)
            self.assertIn("selected_plan", m3)
            self.assertIn("plan_table", m3["selected_plan"])
            pt = m3["selected_plan"]["plan_table"]
            self.assertIn("headers", pt)
            self.assertIn("rows", pt)
            dims = m3["plan_mode"]["dimensions"]
            self.assertIn("timeline", dims)
            self.assertIn("content", dims)
            self.assertIn("metrics", dims)
            print("Report generated successfully.")
        except Exception as e:
            self.fail(f"流程测试异常: {e}")

    def test_03_file_upload_docx(self):
        """测试：上传 DOCX 文件解析"""
        print("\n--- Test 03: File Upload (DOCX) ---")
        
        # 创建临时文件
        filename = "temp_test_resume.docx"
        try:
            doc = Document()
            doc.add_heading('简历', 0)
            doc.add_paragraph('姓名：测试李四')
            doc.add_paragraph('学历：硕士')
            doc.add_paragraph('技能：深度学习, PyTorch, TensorFlow')
            doc.save(filename)
            
            with open(filename, "rb") as f:
                resp = requests.post(
                    f"{BASE_URL}/api/v1/student/parse",
                    files={"file": (filename, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                    timeout=300 # 增加到 5 分钟
                )
            
            self.assertEqual(resp.status_code, 200, f"文件上传解析失败: {resp.text}")
            data = resp.json()
            self.assertEqual(data["student_profile"]["basic_info"]["name"], "测试李四")
            print("DOCX upload test passed.")
            
        except Exception as e:
            self.fail(f"文件上传测试异常: {e}")
        finally:
            # 清理临时文件
            if os.path.exists(filename):
                try:
                    os.remove(filename)
                except PermissionError:
                    pass

    def test_04_invalid_job_id(self):
        """测试：请求不存在的岗位 ID"""
        print("\n--- Test 04: Invalid Job ID ---")
        
        # 构造一个假的 Profile (简化版)
        mock_profile = {
            "basic_info": {"name": "Test", "university": "X", "major": "X", "education_level": "本科"},
            "career_intent": {"expected_salary": "10k", "job_preferences": [], "target_cities": []},
            "competencies": {
                "professional_skills": {"score": 5, "evidence": "", "keywords": []},
                "certificate_requirements": {"score": 5, "evidence": "", "items": [], "missing": []},
                "innovation_capability": {"score": 5, "evidence": ""},
                "learning_capability": {"score": 5, "evidence": ""},
                "stress_resistance": {"score": 5, "evidence": ""},
                "communication_skills": {"score": 5, "evidence": ""},
                "internship_experience": {"score": 5, "history": [], "evaluation": ""}
            },
            "experiences": {"projects": [], "awards": []}
        }
        
        try:
            resp = requests.post(
                f"{BASE_URL}/api/v1/match/report",
                json={
                    "student_profile": mock_profile,
                    "job_id": "INVALID_JOB_ID_999"
                },
                timeout=30
            )
            self.assertEqual(resp.status_code, 404)
            print("Correctly returned 404 for invalid job ID.")
        except Exception as e:
            self.fail(f"异常测试失败: {e}")

if __name__ == "__main__":
    unittest.main()
